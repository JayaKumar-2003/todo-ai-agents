import os
# pyrefly: ignore [missing-import]
from docx import Document

def get_docx_path(session_id: str) -> str:
    """Helper to generate a dynamic filename based on the session ID."""
    return f"todos_{session_id}.docx"

def initialize_docx(session_id: str) -> str:
    """Creates the Word document with a title if it doesn't exist."""
    file_path = get_docx_path(session_id)
    if not os.path.exists(file_path):
        doc = Document()
        doc.add_heading('My TODO List', level=0)
        doc.add_paragraph(f'Created and managed autonomously by the TODO Agent for Session: {session_id}')
        doc.save(file_path)
    return file_path

def add_todo_to_docx(task_title: str, session_id: str) -> str:
    """Appends a task to the session's document as a bullet point."""
    file_path = initialize_docx(session_id)
    doc = Document(file_path)
    
    # Add task as a bullet list item
    doc.add_paragraph(task_title, style='List Bullet')
    doc.save(file_path)
    return f"Successfully added '{task_title}' to the session Word document."

def read_todos_from_docx(session_id: str) -> list:
    """Reads all bullet points from the session's Word document."""
    file_path = initialize_docx(session_id)
    doc = Document(file_path)
    
    # Extract only the bullet points (skipping title and description)
    tasks = []
    for paragraph in doc.paragraphs:
        if paragraph.style.name == 'List Bullet':
            tasks.append(paragraph.text)
    return tasks

